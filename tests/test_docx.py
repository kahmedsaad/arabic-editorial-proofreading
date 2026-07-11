from io import BytesIO

from docx import Document

from app.parsing.docx_intake import extract_text_from_docx_bytes, extract_text_from_upload


def test_docx_extract_arabic():
    doc = Document()
    doc.add_paragraph("عنوان تجريبي")
    doc.add_paragraph("هذه فقرة عربية للتجربة.")
    buffer = BytesIO()
    doc.save(buffer)
    text = extract_text_from_docx_bytes(buffer.getvalue())
    assert "عنوان تجريبي" in text
    assert "فقرة عربية" in text


def test_upload_txt_and_html():
    assert "مرحبا" in extract_text_from_upload("a.txt", "مرحبا".encode())
    html = "<p>مرحبا <b>بالعالم</b></p>".encode()
    text = extract_text_from_upload("a.html", html)
    assert "مرحبا" in text
    assert "<" not in text
